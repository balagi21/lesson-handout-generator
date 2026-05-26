from io import BytesIO
from fastapi import UploadFile
from docx import Document


async def parse_docx(file: UploadFile) -> str:
    """
    Извлекает текст из DOCX
    """
    content = await file.read()

    doc = Document(BytesIO(content))

    result_parts = []

    # Проходим по всем элементам документа в порядке их следования
    for element in doc.element.body:
        # Параграфы
        if element.tag.endswith('p'):
            paragraph = next((p for p in doc.paragraphs if p._element is element), None)
            if paragraph and paragraph.text.strip():
                result_parts.append(paragraph.text)
        # Таблицы
        elif element.tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element is element), None)
            if table:
                # Форматируем таблицу как Markdown для читаемости
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    rows.append('| ' + ' | '.join(cells) + ' |')
                if rows:
                    # Добавляем заголовок таблицы (разделитель)
                    col_count = len(table.rows[0].cells) if table.rows else 1
                    separator = '|' + ' --- |' * col_count
                    result_parts.append('\n'.join([rows[0], separator] + rows[1:]))

    return '\n\n'.join(result_parts)
